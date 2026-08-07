"""文档管理 API"""
import os
import uuid
import shutil
import tempfile
import zipfile
import io
from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import Optional
from app.core.database import get_db
from app.core.auth import get_current_user
from app.core.config import settings
from app.models.document import Document
from app.models.user import User
from app.services.document_converter import pdf_to_word, word_to_pdf
from app.services.document_imager import document_to_images

router = APIRouter()


# ---------- Pydantic schemas ----------

class DocumentUpdate(BaseModel):
    original_name: Optional[str] = None
    doc_category: Optional[str] = None
    case_id: Optional[int] = None
    client_id: Optional[int] = None
    version: Optional[str] = None
    author: Optional[str] = None
    notes: Optional[str] = None
    is_draft: Optional[int] = None
    template_id: Optional[int] = None
    editor_content: Optional[str] = None


class DocumentCreate(BaseModel):
    """纯文本草稿创建（无文件上传）"""
    original_name: str
    doc_category: str = "other"
    case_id: Optional[int] = None
    client_id: Optional[int] = None
    notes: str = ""
    is_draft: int = 1
    template_id: Optional[int] = None
    editor_content: str = ""


class ExportRequest(BaseModel):
    format: str  # "docx" 或 "pdf"
    html: str    # Quill HTML 内容
    title: str = "文档"


# ---------- Helpers ----------

def _doc_to_dict(doc: Document) -> dict:
    return {
        "id": doc.id,
        "filename": doc.filename,
        "original_name": doc.original_name,
        "file_path": doc.file_path,
        "file_size": doc.file_size,
        "file_type": doc.file_type,
        "doc_category": doc.doc_category,
        "case_id": doc.case_id,
        "client_id": doc.client_id,
        "case_number": doc.case.case_number if doc.case else None,
        "client_name": doc.client.name if doc.client else None,
        "version": doc.version,
        "author": doc.author,
        "notes": doc.notes,
        "is_draft": doc.is_draft or 0,
        "template_id": doc.template_id,
        "editor_content": doc.editor_content or "",
        "uploaded_at": doc.uploaded_at.isoformat() if doc.uploaded_at else None,
    }


def _format_size(size_bytes: int) -> str:
    """将字节数转为可读的文件大小"""
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    else:
        return f"{size_bytes / (1024 * 1024):.2f} MB"


# ---------- Routes ----------

@router.get("")
def list_documents(
    search: str = Query("", description="搜索文件名"),
    case_id: Optional[int] = Query(None),
    client_id: Optional[int] = Query(None),
    category: str = Query(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取文档列表"""
    q = db.query(Document)

    if search:
        q = q.filter(Document.original_name.contains(search))
    if case_id:
        q = q.filter(Document.case_id == case_id)
    if client_id:
        q = q.filter(Document.client_id == client_id)
    if category:
        q = q.filter(Document.doc_category == category)

    docs = q.order_by(Document.uploaded_at.desc()).all()
    result = []
    for d in docs:
        item = _doc_to_dict(d)
        item["size_display"] = _format_size(d.file_size)
        result.append(item)
    return result


@router.get("/categories")
def get_doc_categories():
    """获取文档分类列表"""
    return [
        {"value": "legal_opinion", "label": "法律意见书"},
        {"value": "contract_draft", "label": "合同草案"},
        {"value": "complaint", "label": "起诉状"},
        {"value": "defense", "label": "答辩状"},
        {"value": "proxy_statement", "label": "代理词"},
        {"value": "evidence_list", "label": "证据清单"},
        {"value": "other", "label": "其他"},
    ]


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    doc_category: str = Form("other"),
    case_id: Optional[int] = Form(None),
    client_id: Optional[int] = Form(None),
    notes: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """上传文档文件"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="未选择文件")

    # 生成唯一文件名以避重复
    ext = os.path.splitext(file.filename)[1]
    unique_name = f"{uuid.uuid4().hex}{ext}"
    save_path = os.path.join(settings.UPLOAD_DIR, unique_name)

    # 保存文件
    try:
        with open(save_path, "wb") as f:
            content = await file.read()
            f.write(content)
    except Exception:
        raise HTTPException(status_code=500, detail="文件保存失败")

    file_size = os.path.getsize(save_path)

    # 记录到数据库
    doc = Document(
        filename=unique_name,
        original_name=file.filename,
        file_path=save_path,
        file_size=file_size,
        file_type=ext.lstrip("."),
        doc_category=doc_category,
        case_id=case_id if case_id and case_id > 0 else None,
        client_id=client_id if client_id and client_id > 0 else None,
        version=1,
        author=current_user.real_name or current_user.username,
        notes=notes,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    result = _doc_to_dict(doc)
    result["size_display"] = _format_size(doc.file_size)
    return result


@router.get("/{doc_id}/download")
def download_document(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """下载文档文件"""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    if not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="文件已被删除或移动")

    return FileResponse(
        path=doc.file_path,
        filename=doc.original_name,
        media_type="application/octet-stream",
    )


@router.put("/{doc_id}")
def update_document(
    doc_id: int,
    data: DocumentUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """更新文档元数据（JSON body）"""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    update_data = data.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        if key == "case_id":
            setattr(doc, key, value if value and value > 0 else None)
        else:
            setattr(doc, key, value)

    db.commit()
    db.refresh(doc)
    return _doc_to_dict(doc)


@router.post("")
def create_document(
    data: DocumentCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """创建纯文本文档/草稿（无文件上传）"""
    doc = Document(
        filename=f"draft_{uuid.uuid4().hex[:8]}",
        original_name=data.original_name,
        file_path="",  # 草稿无物理文件，用空字符串替代NULL（兼容旧表NOT NULL约束）
        file_size=0,
        file_type="",
        doc_category=data.doc_category,
        case_id=data.case_id if data.case_id and data.case_id > 0 else None,
        client_id=data.client_id if data.client_id and data.client_id > 0 else None,
        version=1,
        author=current_user.real_name or current_user.username,
        notes=data.notes,
        is_draft=data.is_draft,
        template_id=data.template_id if data.template_id and data.template_id > 0 else None,
        editor_content=data.editor_content,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return _doc_to_dict(doc)


@router.delete("/{doc_id}")
def delete_document(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """删除文档（同时删除物理文件）"""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 删除物理文件（草稿无物理文件）
    if doc.file_path and os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except OSError:
            pass

    db.delete(doc)
    db.commit()
    return {"message": "文档删除成功"}


# ==================== 文档格式转换 ====================

@router.post("/convert")
async def convert_document(
    file: UploadFile = File(...),
    target_format: str = Form(...),  # "pdf" 或 "docx"
    current_user: User = Depends(get_current_user),
):
    """将上传的文件在 PDF 和 Word 之间互转

    - target_format="pdf": Word → PDF
    - target_format="docx": PDF → Word
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="未选择文件")

    # 校验目标格式
    if target_format not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="目标格式只支持 pdf 或 docx")

    # 保存临时文件
    ext = os.path.splitext(file.filename)[1].lower()
    tmp_dir = os.path.join(settings.UPLOAD_DIR, "temp")
    os.makedirs(tmp_dir, exist_ok=True)
    tmp_path = os.path.join(tmp_dir, f"{uuid.uuid4().hex}{ext}")

    try:
        with open(tmp_path, "wb") as f:
            content = await file.read()
            f.write(content)
    except Exception:
        raise HTTPException(status_code=500, detail="文件保存失败")

    # 执行转换
    output_path = ""
    if target_format == "pdf":
        output_path = word_to_pdf(tmp_path, tmp_dir)
    elif target_format == "docx":
        output_path = pdf_to_word(tmp_path, tmp_dir)

    # 清理临时文件
    try:
        os.remove(tmp_path)
    except OSError:
        pass

    if not output_path or not os.path.exists(output_path):
        raise HTTPException(status_code=500, detail="文件转换失败，请确认文件格式正确且未损坏")

    output_name = os.path.basename(output_path)
    original_base = os.path.splitext(file.filename)[0]
    download_name = f"{original_base}_转换.{target_format}"

    return FileResponse(
        path=output_path,
        filename=download_name,
        media_type="application/octet-stream",
        headers={"X-Converted-Filename": output_name},
    )


# ==================== 文档图片导出 ====================

@router.post("/export-images")
async def export_document_images(
    file: UploadFile = File(...),
    mode: str = Form("pages"),        # "long" 或 "pages"
    dpi: int = Form(150),
    watermark: str = Form(""),
    current_user: User = Depends(get_current_user),
):
    """将文档导出为图片（长图或分页图）

    - mode=long: 生成一张连续长图
    - mode=pages: 每页一张图（打包为 ZIP 下载）
    - watermark: 可选水印文字
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="未选择文件")

    ext = os.path.splitext(file.filename)[1].lower()
    tmp_dir = os.path.join(settings.UPLOAD_DIR, "temp")
    os.makedirs(tmp_dir, exist_ok=True)
    tmp_path = os.path.join(tmp_dir, f"{uuid.uuid4().hex}{ext}")

    try:
        with open(tmp_path, "wb") as f:
            content = await file.read()
            f.write(content)
    except Exception:
        raise HTTPException(status_code=500, detail="文件保存失败")

    # 导出图片
    file_type = ext.lstrip(".")
    image_paths = document_to_images(
        file_path=tmp_path,
        file_type=file_type,
        output_dir=tmp_dir,
        mode=mode,
        dpi=max(72, min(dpi, 300)),  # 限制 72-300 DPI
        watermark_text=watermark,
    )

    # 清理临时源文件
    try:
        os.remove(tmp_path)
    except OSError:
        pass

    if not image_paths:
        raise HTTPException(status_code=500, detail="图片导出失败，请确认文件格式正确")

    original_base = os.path.splitext(file.filename)[0]

    if mode == "long" and len(image_paths) == 1:
        # 单张长图直接返回
        return FileResponse(
            path=image_paths[0],
            filename=f"{original_base}_长图.png",
            media_type="image/png",
        )
    else:
        # 多张图打包为 ZIP
        zip_path = os.path.join(tmp_dir, f"images_{uuid.uuid4().hex[:8]}.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for i, img_path in enumerate(image_paths):
                arcname = f"{original_base}_第{i+1}页.png"
                zf.write(img_path, arcname)

        # 清理图片文件
        for img_path in image_paths:
            try:
                os.remove(img_path)
            except OSError:
                pass

        return FileResponse(
            path=zip_path,
            filename=f"{original_base}_图片.zip",
            media_type="application/zip",
        )


# ==================== 文件预览 ====================

@router.get("/{doc_id}/preview")
def preview_document(
    doc_id: int,
    page: int = Query(1, ge=1),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """预览文档内容

    - PDF: 直接返回文件流（前端用 PDF.js 渲染）
    - Word: 返回第一页的 PNG 图片
    - 图片: 直接返回文件
    """
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    if not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="文件已被删除或移动")

    file_type = doc.file_type.lower()

    # PDF → 直接流式返回（前端 PDF.js 渲染）
    if file_type == "pdf":
        return FileResponse(
            path=doc.file_path,
            media_type="application/pdf",
        )

    # 图片 → 直接返回
    if file_type in ("png", "jpg", "jpeg", "gif", "bmp", "webp"):
        media_map = {
            "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp",
        }
        return FileResponse(
            path=doc.file_path,
            media_type=media_map.get(file_type, "application/octet-stream"),
        )

    # Word → 转成图片预览
    if file_type in ("doc", "docx", "docm"):
        tmp_dir = os.path.join(settings.UPLOAD_DIR, "temp")
        os.makedirs(tmp_dir, exist_ok=True)
        image_paths = document_to_images(
            file_path=doc.file_path,
            file_type=file_type,
            output_dir=tmp_dir,
            mode="pages",
            dpi=120,
        )
        if not image_paths:
            raise HTTPException(status_code=500, detail="预览生成失败")

        # 返回指定页
        idx = min(page - 1, len(image_paths) - 1)
        preview_path = image_paths[idx]
        return FileResponse(
            path=preview_path,
            media_type="image/png",
        )

    raise HTTPException(status_code=400, detail=f"不支持的文件格式: {file_type}")


# ==================== 文档导出（草稿→Word/PDF） ====================

@router.post("/export")
def export_document(
    data: ExportRequest,
    current_user: User = Depends(get_current_user),
):
    """将 Quill HTML 导出为 Word 或 PDF"""
    if data.format not in ("docx", "pdf"):
        raise HTTPException(status_code=400, detail="格式只支持 docx 或 pdf")

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from bs4 import BeautifulSoup
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"缺少依赖库: {e}")

    tmp_dir = os.path.join(settings.UPLOAD_DIR, "temp")
    os.makedirs(tmp_dir, exist_ok=True)

    doc = DocxDocument()
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 解析 HTML
    soup = BeautifulSoup(data.html, 'html.parser')

    # 标签映射
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ol', 'ul', 'table', 'blockquote']):
        if element.name == 'h1':
            p = doc.add_heading(element.get_text(strip=True), level=1)
        elif element.name == 'h2':
            p = doc.add_heading(element.get_text(strip=True), level=2)
        elif element.name == 'h3':
            p = doc.add_heading(element.get_text(strip=True), level=3)
        elif element.name == 'p':
            p = doc.add_paragraph()
            style_attr = element.get('style', '')
            if 'text-align:right' in style_attr or 'text-align: right' in style_attr:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'text-align:center' in style_attr:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 处理行内元素
            for child in element.children:
                if child.name == 'strong' or child.name == 'b':
                    run = p.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    run = p.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'u':
                    run = p.add_run(child.get_text())
                    run.underline = True
                elif child.string:
                    p.add_run(str(child.string))
        elif element.name in ('ol', 'ul'):
            items = element.find_all('li')
            for li in items:
                p = doc.add_paragraph(li.get_text(strip=True), style='List Bullet' if element.name == 'ul' else 'List Number')
        elif element.name == 'blockquote':
            p = doc.add_paragraph(element.get_text(strip=True))
            p.paragraph_format.left_indent = Cm(1)
            for run in p.runs:
                run.italic = True
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                cols = len(rows[0].find_all(['td', 'th']))
                table = doc.add_table(rows=len(rows), cols=max(cols, 1))
                table.style = 'Table Grid'
                for i, tr in enumerate(rows):
                    cells = tr.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        if j < cols:
                            table.cell(i, j).text = cell.get_text(strip=True)

    # 保存 docx
    docx_path = os.path.join(tmp_dir, f"export_{uuid.uuid4().hex[:8]}.docx")
    doc.save(docx_path)

    if data.format == "docx":
        safe_title = data.title.replace("/", "_").replace("\\", "_")
        return FileResponse(
            path=docx_path,
            filename=f"{safe_title}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        # DOCX → PDF
        pdf_path = word_to_pdf(docx_path, tmp_dir)
        safe_title = data.title.replace("/", "_").replace("\\", "_")
        if pdf_path and os.path.exists(pdf_path):
            return FileResponse(
                path=pdf_path,
                filename=f"{safe_title}.pdf",
                media_type="application/pdf",
            )
        else:
            raise HTTPException(status_code=500, detail="PDF 转换失败，请确认系统已安装 Word 或 LibreOffice")
