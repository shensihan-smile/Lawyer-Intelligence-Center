"""Word ↔ HTML 双向转换服务

提供 .docx 文件与 Quill 编辑器 HTML 之间的双向转换：
- docx_to_html: 将 Word 文档转为可编辑的 HTML 片段
- html_to_docx: 将 Quill HTML 重建为 .docx 文件
"""

from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from bs4 import BeautifulSoup


def docx_to_html(docx_path: str) -> str:
    """将 .docx 文件转为 Quill 可消费的 HTML 片段

    保留：标题级别(h1-h3)、加粗/斜体/下划线、居中对齐、列表、空行。
    表格每行文字以 " | " 连接成文本行（v1 不做表格还原）。
    图片 v1 不提取。

    Args:
        docx_path: .docx 文件路径

    Returns:
        HTML 片段字符串（无 <html>/<body> 包裹）
    """
    doc = DocxDocument(docx_path)
    parts = []

    # 获取段落或表格所属的 heading 级别
    def _heading_level(para):
        style_name = (para.style.name if para.style else "").lower()
        if "heading 1" in style_name:
            return 1
        elif "heading 2" in style_name:
            return 2
        elif "heading 3" in style_name:
            return 3
        return 0

    def _para_to_html(para):
        """将一个 docx 段落转为 HTML"""
        h_level = _heading_level(para)
        align = ""
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align = ' style="text-align:center"'
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align = ' style="text-align:right"'

        # 收集文本和格式
        runs_html = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            # 转义 HTML 特殊字符
            text = (text.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;"))
            if run.bold and run.italic:
                text = f"<strong><em>{text}</em></strong>"
            elif run.bold:
                text = f"<strong>{text}</strong>"
            elif run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            runs_html.append(text)

        inner = "".join(runs_html)

        # 空段落 → Quill 空行
        if not inner.strip():
            inner = "<br>"

        if h_level == 1:
            return f"<h1>{inner}</h1>"
        elif h_level == 2:
            return f"<h2>{inner}</h2>"
        elif h_level == 3:
            return f"<h3>{inner}</h3>"
        else:
            return f"<p{align}>{inner}</p>"

    def _table_to_html(table):
        """将 docx 表格转为文本行（v1 摊平策略）"""
        lines = []
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                # 取单元格所有段落文字
                ct = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                cells_text.append(ct)
            lines.append(" | ".join(cells_text))
        # 返回段落
        inner = "<br>".join(
            line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            for line in lines
        )
        return f"<p>{inner}</p>"

    # 遍历文档 body 子元素，保持段落/表格的先后顺序
    try:
        body = doc.element.body
        for child in body.iterchildren():
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                # 找到对应的 Paragraph 对象
                para = None
                for p in doc.paragraphs:
                    if p._element is child:
                        para = p
                        break
                if para:
                    parts.append(_para_to_html(para))
                else:
                    # 表格内的段落会被 doc.tables 覆盖，这里跳过
                    # 只处理顶层段落
                    pass
            elif tag == "tbl":
                # 找到对应的 Table 对象
                table = None
                for t in doc.tables:
                    if t._element is child:
                        table = t
                        break
                if table:
                    parts.append(_table_to_html(table))
    except Exception:
        # 如果 XML 遍历失败，回退到简单的段落+表格遍历
        # 先输出所有段落
        for para in doc.paragraphs:
            parts.append(_para_to_html(para))
        for table in doc.tables:
            parts.append(_table_to_html(table))

    if not parts:
        return "<p><br></p>"

    return "\n".join(parts)


def html_to_docx(html: str, output_path: str) -> bool:
    """将 Quill HTML 重建为 .docx 文件

    支持：h1/h2/h3、p（含 text-align、strong/em/u）、ol/ul、blockquote、table。
    设置默认字体 SimSun/宋体 12pt。

    Args:
        html: Quill 编辑器 HTML 内容
        output_path: 输出 .docx 文件路径

    Returns:
        成功返回 True，失败返回 False
    """
    try:
        doc = DocxDocument()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "SimSun"
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

        # 解析 HTML
        soup = BeautifulSoup(html, "html.parser")

        for element in soup.find_all(
            ["h1", "h2", "h3", "p", "ol", "ul", "table", "blockquote"]
        ):
            if element.name == "h1":
                doc.add_heading(element.get_text(strip=True), level=1)
            elif element.name == "h2":
                doc.add_heading(element.get_text(strip=True), level=2)
            elif element.name == "h3":
                doc.add_heading(element.get_text(strip=True), level=3)
            elif element.name == "p":
                p = doc.add_paragraph()
                style_attr = element.get("style", "")
                if "text-align:right" in style_attr or "text-align: right" in style_attr:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif "text-align:center" in style_attr:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 处理行内元素
                for child in element.children:
                    if child.name in ("strong", "b"):
                        run = p.add_run(child.get_text())
                        run.bold = True
                    elif child.name in ("em", "i"):
                        run = p.add_run(child.get_text())
                        run.italic = True
                    elif child.name == "u":
                        run = p.add_run(child.get_text())
                        run.underline = True
                    elif child.string:
                        p.add_run(str(child.string))
            elif element.name in ("ol", "ul"):
                items = element.find_all("li")
                for li in items:
                    style_name = (
                        "List Bullet" if element.name == "ul" else "List Number"
                    )
                    p = doc.add_paragraph(li.get_text(strip=True), style=style_name)
            elif element.name == "blockquote":
                p = doc.add_paragraph(element.get_text(strip=True))
                p.paragraph_format.left_indent = Cm(1)
                for run in p.runs:
                    run.italic = True
            elif element.name == "table":
                rows = element.find_all("tr")
                if rows:
                    cols = max(
                        (len(row.find_all(["td", "th"])) for row in rows), default=1
                    )
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = "Table Grid"
                    for i, tr in enumerate(rows):
                        cells = tr.find_all(["td", "th"])
                        for j, cell in enumerate(cells):
                            if j < cols:
                                table.cell(i, j).text = cell.get_text(strip=True)

        doc.save(output_path)
        return True
    except Exception as e:
        print(f"[HTML→DOCX] 转换失败: {e}")
        return False
