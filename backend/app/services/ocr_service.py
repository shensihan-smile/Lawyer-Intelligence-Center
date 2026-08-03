"""OCR 识别服务
支持：
- 图片文件：使用 pytesseract 进行 OCR 识别
- PDF 文件：使用 PyMuPDF 提取文字（原项目已有依赖）
- 如果 pytesseract 未安装，尝试使用系统 Tesseract 命令行
"""
import os
import sys
from pathlib import Path

# 尝试导入 pytesseract
try:
    import pytesseract
    from PIL import Image
    # 指定 Tesseract 安装路径
    _tesseract_paths = [
        r"D:\Tesseract\tesseract.exe",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for _tp in _tesseract_paths:
        if os.path.exists(_tp):
            pytesseract.pytesseract.tesseract_cmd = _tp
            break
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False


def recognize_image(file_path: str) -> str:
    """识别图片中的文字

    Args:
        file_path: 图片文件路径（支持 PNG, JPG, JPEG, BMP, TIFF）

    Returns:
        识别出的文字内容；如果失败返回空字符串
    """
    if not TESSERACT_AVAILABLE:
        return _try_command_line_ocr(file_path)

    try:
        image = Image.open(file_path)
        # 中文 + 英文混合识别
        text = pytesseract.image_to_string(image, lang='chi_sim+eng')
        return text.strip()
    except pytesseract.TesseractNotFoundError:
        # Tesseract 程序未安装
        return "[提示] Tesseract-OCR 程序未安装。请下载安装：https://github.com/UB-Mannheim/tesseract/wiki"
    except Exception as e:
        # 如果中文语言包缺失，尝试只用英文
        try:
            text = pytesseract.image_to_string(Image.open(file_path), lang='eng')
            return text.strip()
        except Exception:
            return f"[识别失败] {str(e)}"


def _try_command_line_ocr(file_path: str) -> str:
    """尝试使用命令行 tesseract（备用方案）"""
    import subprocess
    try:
        output_file = file_path + '_ocr_output'
        result = subprocess.run(
            ['tesseract', file_path, output_file, '-l', 'chi_sim+eng'],
            capture_output=True, text=True, timeout=60
        )
        # 读取输出文件
        txt_file = output_file + '.txt'
        if os.path.exists(txt_file):
            with open(txt_file, 'r', encoding='utf-8') as f:
                text = f.read().strip()
            os.remove(txt_file)
            return text
        return ""
    except FileNotFoundError:
        return "[提示] Tesseract-OCR 未安装。请运行: pip install pytesseract 并安装 Tesseract 程序"
    except Exception as e:
        return f"[命令行OCR失败] {str(e)}"


def extract_pdf_text(file_path: str) -> str:
    """从 PDF 文件中提取文字

    优先使用 PyMuPDF（fitz），原项目已有该依赖

    Args:
        file_path: PDF 文件路径

    Returns:
        提取的文字内容；如果失败返回空字符串
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                text_parts.append(text.strip())
        doc.close()
        return '\n\n'.join(text_parts)
    except ImportError:
        pass
    except Exception as e:
        pass

    # 备用：如果 PDF 是扫描件（图片型PDF），尝试用 OCR
    try:
        return _ocr_pdf_pages(file_path)
    except Exception:
        return ""


def _ocr_pdf_pages(file_path: str) -> str:
    """将 PDF 每页转为图片后 OCR（用于扫描件 PDF）"""
    try:
        from pdf2image import convert_from_path
        images = convert_from_path(file_path, dpi=200, first_page=1, last_page=10)
        text_parts = []
        for i, image in enumerate(images):
            if TESSERACT_AVAILABLE:
                try:
                    text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                except Exception:
                    text = pytesseract.image_to_string(image, lang='eng')
            else:
                # 无 pytesseract，跳过图片OCR
                break
            if text.strip():
                text_parts.append(f"--- 第{i+1}页 ---\n{text.strip()}")
        return '\n\n'.join(text_parts)
    except ImportError:
        return ""
    except Exception as e:
        return f"[PDF图片OCR失败] {str(e)}"


def extract_docx_text(file_path: str) -> str:
    """从 Word .docx 文件中提取文字

    Args:
        file_path: .docx 文件路径

    Returns:
        提取的文字内容；如果失败返回错误提示字符串
    """
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []

        # 提取段落文字
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # 提取表格中的文字
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))

        return '\n\n'.join(text_parts)
    except ImportError:
        return "[提示] python-docx 库未安装，请运行: pip install python-docx"
    except Exception as e:
        return f"[Word文档提取失败] {str(e)}"


def recognize_file(file_path: str, file_type: str = None) -> dict:
    """统一的文件识别入口

    Args:
        file_path: 文件路径
        file_type: 文件类型提示（'image', 'pdf', 或 None 自动检测）

    Returns:
        {
            'text': 识别出的完整文字,
            'summary': 前 200 字摘要,
            'method': 使用的识别方法,
            'page_count': 页数（仅PDF）,
            'success': True/False
        }
    """
    if file_type is None:
        ext = Path(file_path).suffix.lower()
        if ext in ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif', '.gif', '.webp'):
            file_type = 'image'
        elif ext == '.pdf':
            file_type = 'pdf'
        elif ext == '.docx':
            file_type = 'docx'
        else:
            file_type = 'image'  # 默认尝试图片识别

    result = {
        'text': '',
        'summary': '',
        'method': '',
        'page_count': 0,
        'success': False
    }

    if file_type == 'pdf':
        # 先用 PyMuPDF 提取文字
        text = extract_pdf_text(file_path)
        if text.strip():
            result['method'] = 'PyMuPDF 文字提取'
            result['text'] = text
            result['success'] = True
            # 统计页数
            try:
                import fitz
                doc = fitz.open(file_path)
                result['page_count'] = len(doc)
                doc.close()
            except Exception:
                pass
        else:
            # 可能是扫描件 PDF，用 OCR
            result['method'] = 'Tesseract OCR（PDF扫描件）'
            result['text'] = _ocr_pdf_pages(file_path)
            result['success'] = bool(result['text'].strip())
    elif file_type == 'docx':
        # Word 文档
        text = extract_docx_text(file_path)
        result['text'] = text
        result['method'] = 'python-docx 文字提取'
        result['success'] = bool(text.strip()) and not text.startswith('[')
    else:
        # 图片文件
        text = recognize_image(file_path)
        result['text'] = text
        result['method'] = 'Tesseract OCR' if TESSERACT_AVAILABLE else '命令行 Tesseract'
        result['success'] = bool(text.strip()) and not text.startswith('[')

    # 生成摘要（前200字）
    clean_text = result['text'].replace('\n', ' ').replace('\r', ' ').strip()
    result['summary'] = clean_text[:200] if len(clean_text) > 200 else clean_text

    return result


def check_tesseract_available() -> dict:
    """检查 Tesseract 是否可用"""
    status = {
        'pytesseract_installed': TESSERACT_AVAILABLE,
        'tesseract_cli_available': False,
    }

    if TESSERACT_AVAILABLE:
        try:
            version = pytesseract.get_tesseract_version()
            status['tesseract_version'] = str(version)
        except Exception:
            status['tesseract_version'] = 'unknown'

    # 检查命令行 tesseract
    import subprocess
    try:
        result = subprocess.run(['tesseract', '--version'], capture_output=True, text=True, timeout=5)
        status['tesseract_cli_available'] = result.returncode == 0
    except FileNotFoundError:
        pass
    except Exception:
        pass

    return status
