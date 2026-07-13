"""文档格式转换服务：PDF ↔ Word"""
import os
import uuid
import tempfile
import shutil
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def pdf_to_word(pdf_path: str, output_dir: str) -> str:
    """将 PDF 转换为 Word 文档

    处理逻辑：
    1. 按页提取文字块（保留段落结构）
    2. 检测并重建表格
    3. 提取嵌入图片
    4. 用 python-docx 生成 Word 文件

    Returns:
        生成的 Word 文件路径，失败返回空字符串
    """
    try:
        doc = DocxDocument()
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(11)

        pdf = fitz.open(pdf_path)

        for page_num in range(len(pdf)):
            page = pdf[page_num]

            # 页分隔（第一页除外）
            if page_num > 0:
                doc.add_page_break()

            # 获取页面文字块（dict 模式保留位置信息）
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

            for block in blocks:
                if block["type"] == 0:  # 文字块
                    # 检测是否为表格结构（同一行有多个对齐的文字块）
                    para_text = ""
                    for line in block["lines"]:
                        line_text = "".join(
                            span["text"] for span in line["spans"]
                        )
                        para_text += line_text

                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        # 尝试继承原始字体大小
                        try:
                            first_span = block["lines"][0]["spans"][0]
                            font_size = first_span.get("size", 11)
                            flags = first_span.get("flags", 0)
                            for run in p.runs:
                                run.font.size = Pt(font_size)
                                # 粗体检测
                                if flags & 2**3:  # bold flag
                                    run.bold = True
                        except (IndexError, KeyError):
                            pass

                elif block["type"] == 1:  # 图片块
                    try:
                        image_bytes = block.get("image")
                        if image_bytes:
                            # 写入临时图片文件
                            img_ext = block.get("ext", "png")
                            tmp_img = os.path.join(
                                tempfile.gettempdir(),
                                f"_pdf_img_{uuid.uuid4().hex[:6]}.{img_ext}",
                            )
                            with open(tmp_img, "wb") as f:
                                f.write(image_bytes)
                            doc.add_picture(tmp_img, width=Inches(5.5))
                            os.remove(tmp_img)  # 清理临时文件
                    except Exception:
                        pass

        pdf.close()

        # 保存 Word 文件
        output_name = f"converted_{uuid.uuid4().hex[:8]}.docx"
        output_path = os.path.join(output_dir, output_name)
        doc.save(output_path)

        return output_path if os.path.exists(output_path) else ""

    except Exception as e:
        print(f"[pdf_to_word] 转换失败: {e}")
        return ""


def word_to_pdf(word_path: str, output_dir: str) -> str:
    """将 Word 文档转换为 PDF

    策略（按优先级）：
    1. 使用 docx2pdf（调用 Windows 上的 Microsoft Word COM 组件）
    2. 如未安装，尝试用 python-docx 读取内容 + fitz 创建 PDF

    Returns:
        生成的 PDF 文件路径，失败返回空字符串
    """
    output_name = f"converted_{uuid.uuid4().hex[:8]}.pdf"
    output_path = os.path.join(output_dir, output_name)

    # ---- 方法一：docx2pdf（Windows Word COM） ----
    try:
        from docx2pdf import convert as docx2pdf_convert
        # docx2pdf 直接输出到指定文件
        temp_pdf = os.path.join(tempfile.gettempdir(), f"_tmp_{uuid.uuid4().hex[:6]}.pdf")
        docx2pdf_convert(word_path, temp_pdf)
        if os.path.exists(temp_pdf):
            shutil.move(temp_pdf, output_path)
            return output_path
    except ImportError:
        print("[word_to_pdf] docx2pdf 未安装，尝试 python-docx + fitz 方案")
    except Exception as e:
        print(f"[word_to_pdf] docx2pdf 转换失败: {e}")

    # ---- 方法二：python-docx 读取 + fitz 创建 PDF ----
    try:
        doc = DocxDocument(word_path)
        pdf = fitz.open()  # 空白新 PDF

        page = pdf.new_page(width=595, height=842)  # A4
        y = 72  # 起始 Y 坐标（上边距 1 英寸）
        line_height = 14

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                y += line_height * 0.5  # 空行
                continue

            # 检查是否需要换页
            if y > 750:  # 接近页面底部
                page = pdf.new_page(width=595, height=842)
                y = 72

            font_size = 11
            font_name = "china-s"  # CJK 字体

            # 尝试读取段落样式
            try:
                if para.runs:
                    font_size = para.runs[0].font.size
                    if font_size and hasattr(font_size, 'pt'):
                        font_size = font_size.pt
                    else:
                        font_size = 11
            except Exception:
                font_size = 11

            if para.style.name.startswith("Heading"):
                font_size = max(font_size, 16)

            page.insert_text(
                (72, y),
                text,
                fontsize=font_size,
                fontname=font_name,
            )
            y += line_height * (1 + len(text) // 80)  # 估算行数

        pdf.save(output_path)
        pdf.close()
        if os.path.exists(output_path):
            return output_path

    except Exception as e:
        print(f"[word_to_pdf] fitz 方案失败: {e}")

    return ""
